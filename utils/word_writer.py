# Python
import io
import re
from decimal import Decimal
from datetime import date
# python-openxl
from docx import Document
# Django
from django.http import HttpResponse
from django.db.models import Sum
from django.contrib.humanize.templatetags.humanize import intcomma
from django.utils.timezone import localtime
from django.utils import formats
# num2words
from num2words import num2words

def generate_physical_document(request_):
    if request_.regimen_fiscal == 'moral':
        f = open('static/docs/contrato_moral.docx', 'rb')
    else:
        f = open('static/docs/contrato_fisica.docx', 'rb')

    document = Document(f)
    buffer = io.BytesIO()
    places = request_.solicitud_lugar.filter(fecha_reg__year=2024, estatus='assign')
    
    price_places = places.aggregate(price=Sum('precio'))['price'] or 0
    price_extras = places.aggregate(price=Sum('extras__precio'))['price'] or 0
    total_prices = price_places + price_extras
    price_alcohol = places.filter(extras__tipo='licencia_alcohol').aggregate(price=Sum('extras__precio'))['price'] or 0
    alcohol_m2 = places.filter(extras__tipo='licencia_alcohol').aggregate(m2=Sum('extras__m2'))['m2'] or 0
    price_iva = round(total_prices * Decimal(0.16))

    month = formats.date_format(date.today(), 'F')
    day = formats.date_format(date.today(), 'j')
    rfc = ''
    if request_.regimen_fiscal == 'moral' or request_.regimen_fiscal == 'fisica':
        rfc = request_.rfc_txt.upper() if request_.rfc_txt else 'NO ESPECIFICADO'
    else:
        rfc = 'NO APLICA'
    data = {
        '<request_folio>': request_.folio,
        '<razon_social>': request_.nombre.upper(),
        '<address>': '{} {} {} {}'.format(request_.calle, request_.no_calle, request_.codigo_postal, request_.colonia),
        '<town>': request_.municipio,
        '<estate>': request_.get_estado_display(),
        '<rfc>': rfc,
        '<price_no_iva>': intcomma(total_prices-price_iva),
        '<price_no_iva_text>': num2words(total_prices-price_iva, lang='es').upper(),
        '<price_iva>': intcomma(price_iva),
        '<price_iva_text>': num2words(price_iva, lang='es').upper(),
        '<total_iva>': intcomma(total_prices),
        '<total_iva_text>': num2words(total_prices, lang='es').upper(),
        '<day>': day,
        '<month>': month,
        '<name_user>': request_.usuario.get_full_name().upper(),
        '<price_alcohol>': intcomma(price_alcohol),
        '<price_alcohol_text>': num2words(price_alcohol, lang='es'),
        '<m2_alcohol>': str(alcohol_m2),
        '<nombre_replegal>': request_.nombre_replegal.upper() if request_.nombre_replegal else 'NO ESPECIFICADO',
        '<regimen_fiscal>': request_.get_regimen_fiscal_display().upper() if request_.regimen_fiscal else 'NO ESPECIFICADO',
    }
    for p in document.paragraphs:
        if '<table_places>' in p.text:
            table = document.add_table(rows=1, cols=7, style="Table Grid")
            heading_row = table.rows[0].cells
            # add headings
            heading_row[0].text = "Folio"
            heading_row[1].text = "Concepto"
            heading_row[2].text = "m2"
            heading_row[3].text = "Zona"
            heading_row[4].text = "Local"
            heading_row[5].text = "Aplica para"
            heading_row[6].text = "Precio"
            for place in places:
                row = table.add_row().cells
                row[0].text = place.folio
                row[1].text = 'Stand'
                row[2].text = str(place.m2)
                row[3].text = place.get_zona_display()
                row[4].text = place.nombre
                row[6].text =  intcomma(place.precio)
                for px in place.extras.all():
                    row = table.add_row().cells
                    row[0].text = px.folio
                    row[1].text = px.get_tipo_display()
                    row[2].text = str(px.m2)
                    row[5].text = str(px.to_places)
                    row[6].text =  intcomma(px.precio)
            row = table.add_row().cells
            row[5].text = 'Subtotal'
            row[6].text =  str('${}'.format(intcomma(total_prices-price_iva)))
            row = table.add_row().cells
            row[5].text = 'IVA'
            row[6].text =  str('${}'.format(intcomma(price_iva)))
            row = table.add_row().cells
            row[5].text = 'Total'
            row[6].text =  str('${}'.format(intcomma(total_prices)))
            p._p.addnext(table._tbl)
            p.text = ''
        for k, v in data.items():
            regex = re.compile(k)
            paragraph_replace_text(p, regex, v)
    f.close()
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = 'attachment; filename = "CONTRATO_{}.docx"'.format(request_.folio)
    return response

def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph